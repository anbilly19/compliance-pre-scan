import io
from docx import Document
from docx.oxml.ns import qn
from .base import BaseExtractor, ExtractionResult


def _has_vba(doc: Document) -> bool:
    """Check for VBA macros via the vbaProject.bin part."""
    try:
        part_names = [p.partname for p in doc.part.package.iter_parts()]
        return any("vbaProject" in str(p) for p in part_names)
    except Exception:
        return False


def _has_embedded_objects(doc: Document) -> bool:
    """Check for OLE embedded objects in the document body."""
    try:
        body_xml = doc.element.body.xml
        return "oleObject" in body_xml or "embeddedObject" in body_xml
    except Exception:
        return False


class DocxExtractor(BaseExtractor):
    """Extract text from DOCX files."""

    def extract(self, data: bytes, filename: str) -> ExtractionResult:
        doc = Document(io.BytesIO(data))

        paragraphs = [p.text for p in doc.paragraphs]

        # also pull text from tables
        table_cells: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_cells.append(cell.text)

        full_text = "\n".join(paragraphs + table_cells)

        return ExtractionResult(
            text=full_text,
            page_count=0,  # page count not directly available from python-docx
            raw_byte_size=len(data),
            has_macros=_has_vba(doc),
            has_embedded_objects=_has_embedded_objects(doc),
        )
